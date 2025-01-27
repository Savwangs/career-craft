from typing import Optional, List, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from .schemas import Resume
import spacy
from difflib import SequenceMatcher
import re

class ResumeGenerator:
    def __init__(self):
        self.document = Document()
        self.sections = {
            'margins': 1,
            'font_name': 'Calibri',
            'heading_size': 14,
            'subheading_size': 12,
            'body_size': 11
        }
    
    def generate(self, resume: Resume, output_path: str) -> str:
        """Generate a formatted resume document."""
        try:
            # Set up document
            self._setup_document()

            if job_description:
                job_skills = self._extract_job_skills(job_description)
                resume = self._optimize_resume_for_job(resume, job_skills)
            
            # Add sections
            self._add_header(resume)
            self._add_summary(resume)
            self._add_experience(resume)
            self._add_education(resume)
            self._add_skills(resume)
            
            if resume.projects:
                self._add_projects(resume)
            
            if resume.achievements:
                self._add_achievements(resume)
            
            # Save document
            self.document.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Error generating resume: {str(e)}")
            raise

    def _setup_document(self):
        """Set up document formatting."""
        # Set margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(self.sections['margins'])
            section.bottom_margin = Inches(self.sections['margins'])
            section.left_margin = Inches(self.sections['margins'])
            section.right_margin = Inches(self.sections['margins'])

    def _add_header(self, resume: Resume):
        """Add contact information header."""
        # Name
        name = self.document.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name.add_run(resume.contact_info.get('name', 'Name'))
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        
        # Contact info
        contact = self.document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = []
        
        if 'email' in resume.contact_info:
            contact_text.append(resume.contact_info['email'])
        if 'phone' in resume.contact_info:
            contact_text.append(resume.contact_info['phone'])
        if 'location' in resume.contact_info:
            contact_text.append(resume.contact_info['location'])
            
        contact_run = contact.add_run(' | '.join(contact_text))
        contact_run.font.size = Pt(11)
        
        self.document.add_paragraph()  # Add spacing

    def _add_section_heading(self, text: str):
        """Add formatted section heading."""
        heading = self.document.add_paragraph()
        heading_run = heading.add_run(text.upper())
        heading_run.font.size = Pt(self.sections['heading_size'])
        heading_run.font.bold = True
        heading.style = 'Heading 1'

    def _add_summary(self, resume: Resume):
        """Add professional summary section."""
        self._add_section_heading('Professional Summary')
        summary = self.document.add_paragraph()
        summary.add_run(resume.summary).font.size = Pt(self.sections['body_size'])
        self.document.add_paragraph()

    def _add_experience(self, resume: Resume):
        """Add work experience section."""
        self._add_section_heading('Professional Experience')
        
        for exp in resume.experience:
            # Company and position
            p = self.document.add_paragraph()
            company_run = p.add_run(exp.company)
            company_run.font.bold = True
            company_run.font.size = Pt(self.sections['subheading_size'])
            
            # Dates
            dates = self._format_dates(exp.start_date, exp.end_date)
            p.add_run('\t' + dates).font.size = Pt(self.sections['body_size'])
            
            # Position
            position = self.document.add_paragraph()
            position_run = position.add_run(exp.position)
            position_run.font.italic = True
            position_run.font.size = Pt(self.sections['body_size'])
            
            # Description
            if exp.description:
                desc = self.document.add_paragraph()
                desc.add_run(exp.description).font.size = Pt(self.sections['body_size'])
            
            # Highlights
            for highlight in exp.highlights:
                bullet = self.document.add_paragraph(style='List Bullet')
                bullet.add_run(highlight).font.size = Pt(self.sections['body_size'])
            
            self.document.add_paragraph()

    def _add_education(self, resume: Resume):
        """Add education section."""
        self._add_section_heading('Education')
        
        for edu in resume.education:
            p = self.document.add_paragraph()
            
            # Institution
            institution_run = p.add_run(edu.institution)
            institution_run.font.bold = True
            institution_run.font.size = Pt(self.sections['subheading_size'])
            
            # Dates
            dates = self._format_dates(edu.start_date, edu.end_date)
            p.add_run('\t' + dates).font.size = Pt(self.sections['body_size'])
            
            # Degree and field
            degree = self.document.add_paragraph()
            degree_text = f"{edu.degree} in {edu.field_of_study}"
            if edu.gpa:
                degree_text += f" (GPA: {edu.gpa})"
            degree.add_run(degree_text).font.size = Pt(self.sections['body_size'])
            
        self.document.add_paragraph()

    def _add_skills(self, resume: Resume):
        """Add skills section."""
        self._add_section_heading('Skills')
        
        # Group skills by category
        skills_by_category = {}
        for skill in resume.skills:
            if skill.category not in skills_by_category:
                skills_by_category[skill.category] = []
            skills_by_category[skill.category].append(skill.name)
        
        # Add each category
        for category, skills in skills_by_category.items():
            p = self.document.add_paragraph()
            category_run = p.add_run(f"{category}: ")
            category_run.font.bold = True
            category_run.font.size = Pt(self.sections['body_size'])
            
            skills_run = p.add_run(', '.join(skills))
            skills_run.font.size = Pt(self.sections['body_size'])
        
        self.document.add_paragraph()

    def _add_projects(self, resume: Resume):
        """Add projects section."""
        self._add_section_heading('Projects')
        
        for project in resume.projects:
            p = self.document.add_paragraph()
            
            # Project title
            title_run = p.add_run(project.title)
            title_run.font.bold = True
            title_run.font.size = Pt(self.sections['subheading_size'])
            
            # Technologies
            tech_run = p.add_run(f" ({', '.join(project.technologies)})")
            tech_run.font.size = Pt(self.sections['body_size'])
            tech_run.italic = True
            
            # Description
            desc = self.document.add_paragraph()
            desc.add_run(project.description).font.size = Pt(self.sections['body_size'])
            
            if project.url:
                url = self.document.add_paragraph()
                url.add_run(f"URL: {project.url}").font.size = Pt(self.sections['body_size'])
            
            self.document.add_paragraph()

    def _add_achievements(self, resume: Resume):
        """Add achievements section."""
        self._add_section_heading('Achievements')
        
        for achievement in resume.achievements:
            p = self.document.add_paragraph()
            
            # Title
            title_run = p.add_run(achievement.title)
            title_run.font.bold = True
            title_run.font.size = Pt(self.sections['subheading_size'])
            
            # Date if available
            if achievement.date:
                date_str = achievement.date.strftime('%B %Y')
                p.add_run(f" ({date_str})").font.size = Pt(self.sections['body_size'])
            
            # Description
            desc = self.document.add_paragraph()
            desc.add_run(achievement.description).font.size = Pt(self.sections['body_size'])
            
        self.document.add_paragraph()

    def _format_dates(self, start_date: datetime, end_date: Optional[datetime]) -> str:
        """Format date range for display."""
        start_str = start_date.strftime('%B %Y')
        end_str = end_date.strftime('%B %Y') if end_date else 'Present'
        return f"{start_str} - {end_str}"

class ResumeOptimizer:
    def __init__(self):
        # Load spaCy model for natural language processing
        self.nlp = spacy.load('en_core_web_sm')

    def optimize_resume(self, resume: Dict, job_description: str) -> Dict:
        """Optimize resume content while preserving original structure"""
        optimized_resume = resume.copy()

        # Optimize summary
        optimized_resume['summary'] = self._enhance_summary(
            resume['summary'], 
            job_description
        )

        # Optimize experience entries
        optimized_resume['experience'] = [
            self._enhance_experience_entry(exp, job_description) 
            for exp in resume['experience']
        ]

        # Prioritize skills
        optimized_resume['skills'] = self._prioritize_skills(
            resume['skills'], 
            job_description
        )

        return optimized_resume

    def _enhance_summary(self, original_summary: str, job_description: str) -> str:
        """Make summary more targeted and impactful"""
        job_doc = self.nlp(job_description.lower())
        summary_doc = self.nlp(original_summary.lower())

        # Extract key job skills and requirements
        job_keywords = [token.text for token in job_doc 
                        if token.pos_ in ['NOUN', 'VERB'] 
                        and len(token.text) > 3]

        # Rewrite summary incorporating relevant keywords
        enhanced_summary_parts = []
        for sent in summary_doc.sents:
            sent_text = sent.text
            for keyword in job_keywords:
                if keyword in sent_text.lower():
                    sent_text = sent_text.replace(
                        keyword, 
                        f"**{keyword}**"  # Emphasize matching keywords
                    )
            enhanced_summary_parts.append(sent_text)

        return " ".join(enhanced_summary_parts)

    def _enhance_experience_entry(self, experience: Dict, job_description: str) -> Dict:
        """Enhance experience description with more impactful language"""
        original_desc = experience.get('description', '')
        job_doc = self.nlp(job_description.lower())
        desc_doc = self.nlp(original_desc.lower())

        # Extract achievement-oriented verbs from job description
        action_keywords = [
            token.lemma_ for token in job_doc 
            if token.pos_ == 'VERB' and token.lemma_ not in ['be', 'have', 'do']
        ]

        # Rewrite description using more powerful language
        enhanced_desc_parts = []
        for sent in desc_doc.sents:
            sent_text = sent.text
            for verb in action_keywords:
                if verb in sent_text.lower():
                    sent_text = sent_text.replace(
                        verb, 
                        f"**{verb.upper()}**"  # Emphasize powerful action verbs
                    )
            enhanced_desc_parts.append(sent_text)

        experience_copy = experience.copy()
        experience_copy['description'] = " ".join(enhanced_desc_parts)
        return experience_copy

    def _prioritize_skills(self, skills: List[Dict], job_description: str) -> List[Dict]:
        """Reorder skills based on job description relevance"""
        job_doc = self.nlp(job_description.lower())
        
        # Extract skills from job description
        job_skills = [token.text for token in job_doc 
                      if token.pos_ in ['NOUN'] 
                      and len(token.text) > 3]

        def skill_relevance(skill):
            """Calculate skill relevance score"""
            skill_name = skill['name'].lower()
            matches = sum(1 for js in job_skills if js in skill_name)
            return matches

        # Sort skills by relevance to job description
        return sorted(skills, key=skill_relevance, reverse=True)